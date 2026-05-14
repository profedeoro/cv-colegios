"""Orquestador del módulo `generar`.

Para cada colegio en estado 'enriquecido':
1. Lee el perfil pedagógico del colegio.
2. Lee texto del CV pulido (data/cv_base_polished.pdf) — referencia para el validador.
3. Llama a Claude 3 veces: reescribir Perfil, reordenar Bullets, redactar Carta.
4. Cada respuesta pasa por `detectar_alucinaciones`; si encuentra hechos inventados,
   regenera (máx 3 intentos por paso).
5. Genera el PDF personalizado en `data/salida/{slug}.pdf`.
6. Inserta una fila en `borradores` con asunto + cuerpo de carta + ruta del PDF.

Si en cualquier paso se agotan los 3 intentos por alucinación, el colegio se marca
`revisar_manualmente` (transición permitida desde `enriquecido`) y NO se inserta
borrador. El colegio se incrementa en `intentos_generar`. La transición
`enriquecido → borrador_creado` la hace el módulo `enviar_borradores` cuando
Gmail devuelve el draft_id (ver Tarea 7 del Plan 4).
"""
import json
import re
import time
from pathlib import Path
from typing import Callable

from docx import Document

from modulos.db import (
    cambiar_estado,
    colegios_para_generar,
    incrementar_intento_generar,
    insertar_borrador,
    registrar_ejecucion,
)
from modulos.docx_personalizado import generar_pdf_personalizado
from modulos.logger import obtener_logger
from modulos.pdf_lector import leer_pdf
from modulos.validador import detectar_alucinaciones, extraer_hechos


RAIZ = Path(__file__).parent.parent
RUTA_PLANTILLA = RAIZ / "data" / "plantilla_base.docx"
RUTA_POLISHED_PDF = RAIZ / "data" / "cv_base_polished.pdf"
RUTA_POLISHED_DOCX = RAIZ / "data" / "cv_base_polished.docx"
RUTA_SALIDA = RAIZ / "data" / "salida"

RUTA_PROMPT_PERFIL = RAIZ / "prompts" / "reescribir_perfil.txt"
RUTA_PROMPT_BULLETS = RAIZ / "prompts" / "personalizar_bullets.txt"
RUTA_PROMPT_CARTA = RAIZ / "prompts" / "carta_presentacion.txt"

RE_PLACEHOLDER = re.compile(r"\{\{(\w+)\}\}")
RE_EXP_BULLETS = re.compile(r"^EXP_(\d+)_BULLETS$")

MAX_INTENTOS_ALUCINACION = 3
MAX_TOKENS_GENERACION = 1500


# == Extracción del CV base ==

def _extraer_cv_base(plantilla_path: Path | str, polished_docx_path: Path | str) -> dict[str, str]:
    """Mapea cada placeholder de la plantilla al texto correspondiente del polished docx.

    Las dos plantillas comparten estructura (mismo número de párrafos, mismas posiciones).
    Donde la plantilla tiene `{{X}}`, el polished tiene el texto original. Cuando la
    plantilla colapsa varios párrafos en un solo placeholder, los párrafos siguientes
    en la plantilla quedan vacíos; en esos casos, agrupamos el texto del polished hasta
    el próximo párrafo de la plantilla que no sea vacío ni placeholder.

    Devuelve un dict {placeholder: texto}. Los bullets se devuelven como string
    multilínea (un bullet por línea).
    """
    tmpl = Document(str(plantilla_path))
    pol = Document(str(polished_docx_path))
    if len(tmpl.paragraphs) != len(pol.paragraphs):
        raise ValueError(
            f"plantilla ({len(tmpl.paragraphs)} pp) y polished "
            f"({len(pol.paragraphs)} pp) tienen distinto número de párrafos"
        )

    valores: dict[str, str] = {}
    n = len(tmpl.paragraphs)
    i = 0
    while i < n:
        tt = tmpl.paragraphs[i].text
        m = RE_PLACEHOLDER.search(tt)
        if not m:
            i += 1
            continue
        nombre = m.group(1)
        # Recolectar líneas del polished desde i hasta el próximo párrafo no vacío
        # de la plantilla (sea otro placeholder o texto fijo).
        lineas = [pol.paragraphs[i].text]
        j = i + 1
        while j < n and tmpl.paragraphs[j].text.strip() == "":
            texto_polished = pol.paragraphs[j].text
            if texto_polished.strip():
                lineas.append(texto_polished)
            j += 1
        valores[nombre] = "\n".join(lineas).strip()
        i = j
    return valores


def _texto_cv_completo(polished_pdf_path: Path | str) -> str:
    """Devuelve el texto plano del PDF pulido (referencia anti-alucinación)."""
    return leer_pdf(polished_pdf_path)


# == Validador anti-alucinación con retry ==

def _intentar_con_validacion(
    generador: Callable[[], tuple[str, float]],
    *,
    cv_original: str,
    nombres_permitidos: set[str],
    max_intentos: int = MAX_INTENTOS_ALUCINACION,
) -> tuple[str, float, int] | None:
    """Ejecuta `generador()` hasta `max_intentos` veces. Acepta el primer texto sin alucinaciones.

    Devuelve (texto, costo_acumulado_usd, intentos_usados) o None si todos los intentos
    fueron rechazados por el validador.
    """
    costo_total = 0.0
    for intento in range(1, max_intentos + 1):
        texto, costo = generador()
        costo_total += costo
        alucinados = detectar_alucinaciones(cv_original, texto, nombres_permitidos)
        if not alucinados:
            return texto, costo_total, intento
    return None


# == Generadores por paso (perfil / bullets / carta) ==

def _mensaje_perfil(perfil_actual: str, perfil_colegio: dict, nombre_colegio: str) -> str:
    return (
        f"PERFIL_ACTUAL:\n{perfil_actual}\n\n"
        f"PERFIL_COLEGIO:\n{json.dumps(perfil_colegio, ensure_ascii=False, indent=2)}\n\n"
        f"NOMBRE_COLEGIO: {nombre_colegio}"
    )


def _mensaje_bullets(titulo: str, bullets_actuales: str, perfil_colegio: dict) -> str:
    return (
        f"TITULO_EXPERIENCIA: {titulo}\n\n"
        f"BULLETS_ACTUALES:\n{bullets_actuales}\n\n"
        f"PERFIL_COLEGIO:\n{json.dumps(perfil_colegio, ensure_ascii=False, indent=2)}"
    )


def _mensaje_carta(
    nombre_colegio: str, ciudad: str, perfil_colegio: dict, resumen_daniel: str
) -> str:
    return (
        f"NOMBRE_COLEGIO: {nombre_colegio}\n"
        f"CIUDAD: {ciudad}\n\n"
        f"PERFIL_COLEGIO:\n{json.dumps(perfil_colegio, ensure_ascii=False, indent=2)}\n\n"
        f"RESUMEN_DANIEL:\n{resumen_daniel}"
    )


def _llamar_claude(cliente_claude, sistema_path: Path, usuario: str) -> tuple[str, float]:
    sistema = sistema_path.read_text(encoding="utf-8")
    return cliente_claude.preguntar(
        sistema=sistema,
        usuario=usuario,
        max_tokens=MAX_TOKENS_GENERACION,
        cachear_sistema=True,
    )


# == Helpers internos ==

def _claves_experiencias(valores: dict[str, str]) -> list[int]:
    """Devuelve la lista de N para los EXP_N presentes en `valores`, ordenados."""
    ns = set()
    for clave in valores:
        m = RE_EXP_BULLETS.match(clave)
        if m:
            ns.add(int(m.group(1)))
    return sorted(ns)


def _asunto(nombre_colegio: str) -> str:
    """Construye el asunto según spec 4.3."""
    return f"Postulación docente — Daniel E. Villalba — {nombre_colegio}"


def _nombres_permitidos(colegio: dict) -> set[str]:
    """Tokens que el validador debe ignorar al detectar alucinaciones.

    El validador tokeniza nombres propios y números, así que necesitamos pasarle
    los TOKENS del nombre del colegio + ciudad (no los strings completos). Por
    ejemplo, para "Colegio Bilingüe San José" extrae {"Colegio", "Bilingüe",
    "San José"} y todos esos deben permitirse cuando aparezcan en la carta.
    """
    permitidos: set[str] = set()
    nombre = colegio.get("nombre", "") or ""
    ciudad = colegio.get("ciudad", "") or ""
    # Empezamos el texto con palabras minúsculas para que `extraer_hechos`
    # capture la primera palabra mayúscula del nombre como hecho (sin esa
    # antesala, las palabras al inicio de oración se ignoran).
    texto_referencia = f"el colegio es {nombre} en {ciudad}."
    permitidos.update(extraer_hechos(texto_referencia))
    # También el string completo, por si el extractor lo encuentra entero.
    if nombre:
        permitidos.add(nombre)
    if ciudad:
        permitidos.add(ciudad)
    return permitidos


# == Orquestador por colegio ==

def _generar_perfil_validado(
    cliente_claude, perfil_actual: str, perfil_colegio: dict,
    nombre_colegio: str, cv_original: str, permitidos: set[str],
) -> tuple[str, float, int] | None:
    """Reescribe el bloque Perfil. Devuelve (texto, costo, intentos) o None si se rinde."""
    return _intentar_con_validacion(
        lambda: _llamar_claude(
            cliente_claude, RUTA_PROMPT_PERFIL,
            _mensaje_perfil(perfil_actual, perfil_colegio, nombre_colegio),
        ),
        cv_original=cv_original, nombres_permitidos=permitidos,
    )


def _generar_bullets_validado(
    cliente_claude, titulo: str, bullets_actuales: str, perfil_colegio: dict,
    cv_original: str, permitidos: set[str],
) -> tuple[str, float, int] | None:
    """Reordena/refrasea bullets de una experiencia."""
    return _intentar_con_validacion(
        lambda: _llamar_claude(
            cliente_claude, RUTA_PROMPT_BULLETS,
            _mensaje_bullets(titulo, bullets_actuales, perfil_colegio),
        ),
        cv_original=cv_original, nombres_permitidos=permitidos,
    )


def _generar_carta_validada(
    cliente_claude, nombre_colegio: str, ciudad: str, perfil_colegio: dict,
    resumen_daniel: str, cv_original: str, permitidos: set[str],
) -> tuple[str, float, int] | None:
    """Redacta el cuerpo de la carta de presentación."""
    return _intentar_con_validacion(
        lambda: _llamar_claude(
            cliente_claude, RUTA_PROMPT_CARTA,
            _mensaje_carta(nombre_colegio, ciudad, perfil_colegio, resumen_daniel),
        ),
        cv_original=cv_original, nombres_permitidos=permitidos,
    )


def procesar_colegio(
    ruta_bd,
    colegio: dict,
    cliente_claude,
    *,
    plantilla_path: Path = RUTA_PLANTILLA,
    polished_pdf_path: Path = RUTA_POLISHED_PDF,
    polished_docx_path: Path = RUTA_POLISHED_DOCX,
    salida_dir: Path = RUTA_SALIDA,
) -> dict:
    """Genera HV personalizada + carta para un colegio. Inserta borrador si tiene éxito.

    Devuelve dict con resumen del resultado:
        {"colegio_id": int, "estado_final": str, "costo": float, "razon": str|None}
    `estado_final` ∈ {"borrador_insertado", "revisar_manualmente"}.
    """
    log = obtener_logger("generar")
    cid = colegio["id"]
    nombre = colegio["nombre"]
    ciudad = colegio["ciudad"]

    perfil_colegio = colegio.get("perfil_pedagogico")
    if isinstance(perfil_colegio, str):
        perfil_colegio = json.loads(perfil_colegio)
    if not perfil_colegio:
        log.warning(f"[{cid}] '{nombre}' sin perfil_pedagogico; se marca para revisión")
        incrementar_intento_generar(ruta_bd, cid)
        cambiar_estado(ruta_bd, cid, "revisar_manualmente")
        return {"colegio_id": cid, "estado_final": "revisar_manualmente",
                "costo": 0.0, "razon": "sin_perfil_pedagogico"}

    valores_base = _extraer_cv_base(plantilla_path, polished_docx_path)
    cv_original = _texto_cv_completo(polished_pdf_path)
    permitidos = _nombres_permitidos(colegio)
    costo_total = 0.0

    # Paso 1: reescribir Perfil
    res = _generar_perfil_validado(
        cliente_claude, valores_base["PERFIL"], perfil_colegio, nombre,
        cv_original, permitidos,
    )
    if res is None:
        return _marcar_giveup(ruta_bd, cid, log, "perfil", costo_total)
    texto_perfil, costo, _ = res
    costo_total += costo
    valores_base["PERFIL"] = texto_perfil

    # Paso 2: reordenar bullets de cada experiencia
    for n in _claves_experiencias(valores_base):
        res = _generar_bullets_validado(
            cliente_claude, valores_base[f"EXP_{n}_TITULO"],
            valores_base[f"EXP_{n}_BULLETS"], perfil_colegio,
            cv_original, permitidos,
        )
        if res is None:
            return _marcar_giveup(ruta_bd, cid, log, f"bullets EXP_{n}", costo_total)
        texto_bullets, costo, _ = res
        costo_total += costo
        valores_base[f"EXP_{n}_BULLETS"] = texto_bullets

    # Paso 3: redactar carta de presentación (usando perfil ya reescrito como resumen)
    res = _generar_carta_validada(
        cliente_claude, nombre, ciudad, perfil_colegio, valores_base["PERFIL"],
        cv_original, permitidos,
    )
    if res is None:
        return _marcar_giveup(ruta_bd, cid, log, "carta", costo_total)
    texto_carta, costo, _ = res
    costo_total += costo

    # Generar PDF personalizado e insertar borrador
    slug = colegio.get("nombre_normalizado") or str(cid)
    ruta_pdf = Path(salida_dir) / f"{slug}.pdf"
    ruta_pdf.parent.mkdir(parents=True, exist_ok=True)
    generar_pdf_personalizado(plantilla_path, valores_base, ruta_pdf)
    insertar_borrador(
        ruta_bd, cid,
        tipo="inicial",
        asunto=_asunto(nombre),
        cuerpo_carta=texto_carta,
        ruta_pdf_hv=str(ruta_pdf),
    )
    log.info(f"[{cid}] Borrador insertado para '{nombre}' (costo: ${costo_total:.4f})")
    return {
        "colegio_id": cid,
        "estado_final": "borrador_insertado",
        "costo": costo_total,
        "razon": None,
    }


def _marcar_giveup(ruta_bd, cid: int, log, paso: str, costo: float) -> dict:
    """Después de 3 intentos fallidos por alucinación: incrementa intento y marca para revisión."""
    incrementar_intento_generar(ruta_bd, cid)
    cambiar_estado(ruta_bd, cid, "revisar_manualmente")
    log.warning(
        f"[{cid}] Alucinación recurrente en paso '{paso}'. Marcado revisar_manualmente. "
        f"Costo desperdiciado: ${costo:.4f}"
    )
    return {
        "colegio_id": cid,
        "estado_final": "revisar_manualmente",
        "costo": costo,
        "razon": f"alucinacion_{paso}",
    }


# == Orquestador del módulo completo ==

def ejecutar(
    ruta_bd,
    cliente_claude,
    max_colegios: int = 15,
    *,
    plantilla_path: Path = RUTA_PLANTILLA,
    polished_pdf_path: Path = RUTA_POLISHED_PDF,
    polished_docx_path: Path = RUTA_POLISHED_DOCX,
    salida_dir: Path = RUTA_SALIDA,
) -> dict:
    """Procesa hasta `max_colegios` enriquecidos y devuelve resumen."""
    log = obtener_logger("generar")
    log.info(f"Iniciando generación (max={max_colegios})")
    inicio = time.monotonic()

    pendientes = colegios_para_generar(ruta_bd, limite=max_colegios)
    log.info(f"Pendientes encontrados: {len(pendientes)}")

    resumen: dict[str, int] = {}
    costo_total = 0.0
    for col in pendientes:
        try:
            resultado = procesar_colegio(
                ruta_bd, col, cliente_claude,
                plantilla_path=plantilla_path,
                polished_pdf_path=polished_pdf_path,
                polished_docx_path=polished_docx_path,
                salida_dir=salida_dir,
            )
        except Exception as e:
            log.exception(f"[{col['id']}] Error inesperado generando: {e}")
            incrementar_intento_generar(ruta_bd, col["id"])
            resumen["error"] = resumen.get("error", 0) + 1
            continue
        clave = resultado["estado_final"]
        resumen[clave] = resumen.get(clave, 0) + 1
        costo_total += resultado.get("costo", 0.0)

    duracion = time.monotonic() - inicio
    estado_ejec = "ok" if resumen.get("error", 0) == 0 else "error"
    registrar_ejecucion(
        ruta_bd, modulo="generar", duracion_segundos=duracion,
        estado=estado_ejec, colegios_procesados=len(pendientes),
        costo_api_usd=costo_total,
    )
    log.info(
        f"Generación terminada en {duracion:.1f}s. Resumen: {resumen}. "
        f"Costo total: ${costo_total:.4f}"
    )
    return {"resumen": resumen, "costo_usd": costo_total, "duracion_seg": duracion}
