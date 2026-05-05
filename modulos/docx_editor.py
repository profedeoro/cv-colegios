"""Editor de DOCX que preserva estilos.

Este módulo difiere de `modulos.plantilla` en que opera sobre documentos
con formato existente (CV del usuario), preservando estilos a nivel de run
cuando se hacen correcciones de typos. La función `reemplazar_texto_parrafo`
sí destruye estilos sub-párrafo, pero preserva el estilo del primer run
(útil para insertar placeholders en párrafos completos).
"""
from docx.document import Document as DocumentType


def enumerar_parrafos(doc: DocumentType) -> list[tuple[int, str]]:
    """Devuelve [(idx, texto)] solo para párrafos con texto no vacío."""
    pares = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            pares.append((i, p.text))
    return pares


def aplicar_correcciones_a_parrafo(parrafo, correcciones: list[tuple[str, str]]) -> int:
    """Aplica correcciones (find, replace) a los runs del párrafo. Devuelve total de reemplazos."""
    total = 0
    for find, replace in correcciones:
        for run in parrafo.runs:
            if find in run.text:
                ocurrencias = run.text.count(find)
                run.text = run.text.replace(find, replace)
                total += ocurrencias
    return total


def aplicar_correcciones_a_documento(doc: DocumentType, correcciones: list[tuple[str, str]]) -> int:
    """Aplica correcciones a TODOS los párrafos del documento, incluidas tablas."""
    total = 0
    for parrafo in doc.paragraphs:
        total += aplicar_correcciones_a_parrafo(parrafo, correcciones)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    total += aplicar_correcciones_a_parrafo(parrafo, correcciones)
    return total


def reemplazar_texto_parrafo(parrafo, nuevo_texto: str) -> None:
    """Reemplaza todo el texto del párrafo poniéndolo en runs[0], limpiando los demás.

    Preserva el estilo (fuente, negrilla, etc.) del primer run.
    Si el párrafo no tiene runs, agrega uno nuevo.
    """
    if not parrafo.runs:
        parrafo.add_run(nuevo_texto)
        return
    parrafo.runs[0].text = nuevo_texto
    for run in parrafo.runs[1:]:
        run.text = ""
