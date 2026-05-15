import re
from pathlib import Path
from docx import Document

RE_PLACEHOLDER = re.compile(r"\{\{(\w+)\}\}")


def _set_parrafo_texto(parrafo, texto):
    """Reemplaza el texto del párrafo preservando el formato del primer run."""
    for run in parrafo.runs:
        run.text = ""
    if parrafo.runs:
        parrafo.runs[0].text = texto
    else:
        parrafo.add_run(texto)


def rellenar_plantilla(
    plantilla_path: Path | str,
    salida_path: Path | str,
    valores: dict[str, str],
) -> None:
    """Carga una plantilla DOCX, reemplaza {{NOMBRE}}-style placeholders, guarda en salida.

    Si el valor a substituir es multi-línea (contiene `\\n`), las líneas se distribuyen
    a través del párrafo del placeholder + los párrafos vacíos consecutivos que le
    siguen. Si hay más líneas que párrafos disponibles, las líneas sobrantes se
    concatenan en el último párrafo (separadas por `\\n`).
    """
    doc = Document(str(plantilla_path))
    placeholders_encontrados = set()
    parrafos = list(doc.paragraphs)

    i = 0
    while i < len(parrafos):
        p = parrafos[i]
        texto = p.text
        # Track which placeholders we see
        for m in RE_PLACEHOLDER.finditer(texto):
            placeholders_encontrados.add(m.group(1))
        # Substitute placeholders in this paragraph's text
        nuevo = RE_PLACEHOLDER.sub(
            lambda m: valores.get(m.group(1), m.group(0)),
            texto,
        )
        if nuevo == texto:
            i += 1
            continue
        # Multi-line case: distribute across this para + adjacent empties
        if "\n" in nuevo:
            lineas = nuevo.split("\n")
            # Find consecutive empty paragraphs ahead
            empties = []
            j = i + 1
            while j < len(parrafos) and not parrafos[j].text.strip():
                empties.append(parrafos[j])
                j += 1
            slots = [p] + empties
            if len(lineas) > len(slots):
                # Group excess lines into the last slot
                head = lineas[:len(slots) - 1]
                tail = "\n".join(lineas[len(slots) - 1:])
                lineas = head + [tail]
            for slot, linea in zip(slots, lineas):
                _set_parrafo_texto(slot, linea)
            i = j  # skip the empties we just filled
        else:
            _set_parrafo_texto(p, nuevo)
            i += 1

    # Tables: same logic without empty-sibling distribution (cells are atomic)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    texto = parrafo.text
                    for m in RE_PLACEHOLDER.finditer(texto):
                        placeholders_encontrados.add(m.group(1))
                    nuevo = RE_PLACEHOLDER.sub(
                        lambda m: valores.get(m.group(1), m.group(0)),
                        texto,
                    )
                    if nuevo != texto:
                        _set_parrafo_texto(parrafo, nuevo)

    faltantes = placeholders_encontrados - set(valores.keys())
    if faltantes:
        raise ValueError(f"placeholder(s) sin valor: {sorted(faltantes)}")

    Path(salida_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(salida_path))
