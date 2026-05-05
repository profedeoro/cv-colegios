"""Pule el CV base (DOCX) con Claude y produce plantilla.docx + cv_pulido.pdf.

Lee `data/cv_base.docx`, manda los párrafos a Claude para identificar correcciones
ortográficas y mapa de placeholders. Aplica los cambios preservando los estilos
originales (fuente, negrillas, alineaciones).

Uso:
    python reconstruir_plantilla.py
"""
import copy
import hashlib
import json
from pathlib import Path
from typing import Callable

from docx import Document

from modulos.cliente_claude import ClienteClaude
from modulos.config import cargar_config
from modulos.db import guardar_hash_cv
from modulos.docx_editor import (
    aplicar_correcciones_a_documento,
    enumerar_parrafos,
    reemplazar_texto_parrafo,
)
from modulos.pdf_conversor import convertir_docx_a_pdf

RAIZ = Path(__file__).parent
RUTA_CV_BASE = RAIZ / "data" / "cv_base.docx"
RUTA_SALIDA_DOCX = RAIZ / "data" / "plantilla_base.docx"
RUTA_SALIDA_DOCX_PULIDO = RAIZ / "data" / "cv_base_polished.docx"
RUTA_SALIDA_PDF = RAIZ / "data" / "cv_base_polished.pdf"
RUTA_PROMPT = RAIZ / "prompts" / "analizar_cv_docx.txt"
RUTA_BD = RAIZ / "data" / "colegios.db"


def _formato_para_claude(pares: list[tuple[int, str]]) -> str:
    """Formatea los párrafos como '[idx] texto' separados por saltos de línea."""
    return "\n".join(f"[{i}] {t}" for i, t in pares)


def _parsear_respuesta_claude(respuesta: str) -> dict:
    """Parsea el JSON, tolerante a code fences."""
    raw = respuesta.strip()
    if raw.startswith("```"):
        raw = raw.split("```", 2)[1]
        if raw.lower().lstrip().startswith("json"):
            if "\n" in raw:
                raw = raw[raw.index("\n") + 1:]
        raw = raw.rsplit("```", 1)[0]
    return json.loads(raw.strip())


def ejecutar(
    cv_docx: Path = RUTA_CV_BASE,
    salida_docx_plantilla: Path = RUTA_SALIDA_DOCX,
    salida_docx_pulido: Path = RUTA_SALIDA_DOCX_PULIDO,
    salida_pdf: Path = RUTA_SALIDA_PDF,
    api_key: str | None = None,
    confirmar: Callable[[dict], bool] | None = None,
    ruta_bd: Path = RUTA_BD,
) -> None:
    """Pule el CV y genera plantilla con placeholders, preservando estilos."""
    if api_key is None:
        config = cargar_config()
        api_key = config["ANTHROPIC_API_KEY"]

    if not cv_docx.exists():
        raise FileNotFoundError(
            f"No se encontró {cv_docx}. Coloca tu HV en formato .docx en esa ruta."
        )

    if not RUTA_PROMPT.exists():
        raise FileNotFoundError(
            f"No se encontró el prompt en {RUTA_PROMPT}. El repositorio puede estar incompleto."
        )

    # 1. Leer DOCX y enumerar párrafos no vacíos
    doc_origen = Document(str(cv_docx))
    pares = enumerar_parrafos(doc_origen)
    if not pares:
        raise ValueError(f"El DOCX {cv_docx} no contiene texto.")
    texto_para_claude = _formato_para_claude(pares)

    # 2. Llamar a Claude con el prompt de análisis
    sistema = RUTA_PROMPT.read_text(encoding="utf-8")
    cliente = ClienteClaude(api_key=api_key)
    respuesta, costo = cliente.preguntar(sistema=sistema, usuario=texto_para_claude, max_tokens=8000)
    datos = _parsear_respuesta_claude(respuesta)

    # 3. Confirmación interactiva (opcional)
    if confirmar is not None and not confirmar(datos):
        print("Cancelado por el usuario. No se guardó nada.")
        return

    # 4. Construir lista de tuplas (find, replace) a partir del JSON
    correcciones = [(c["buscar"], c["reemplazar"]) for c in datos.get("correcciones", [])]

    # 5. Generar el DOCX pulido (copia del original con correcciones aplicadas)
    doc_pulido = Document(str(cv_docx))
    n_pulido = aplicar_correcciones_a_documento(doc_pulido, correcciones)
    salida_docx_pulido.parent.mkdir(parents=True, exist_ok=True)
    doc_pulido.save(str(salida_docx_pulido))

    # 6. Convertir el DOCX pulido a PDF
    convertir_docx_a_pdf(salida_docx_pulido, salida_pdf)

    # 7. Generar la plantilla con placeholders (otra copia + correcciones + reemplazos)
    doc_plantilla = Document(str(cv_docx))
    aplicar_correcciones_a_documento(doc_plantilla, correcciones)

    # Reemplazar bloque Perfil: el primer índice se vuelve {{PERFIL}}, los demás se vacían
    indices_perfil = datos.get("bloque_perfil", [])
    if indices_perfil:
        reemplazar_texto_parrafo(doc_plantilla.paragraphs[indices_perfil[0]], "{{PERFIL}}")
        for idx in indices_perfil[1:]:
            reemplazar_texto_parrafo(doc_plantilla.paragraphs[idx], "")

    # Reemplazar cada experiencia
    for n, exp in enumerate(datos.get("experiencias", []), start=1):
        if "titulo_idx" in exp:
            reemplazar_texto_parrafo(
                doc_plantilla.paragraphs[exp["titulo_idx"]],
                f"{{{{EXP_{n}_TITULO}}}}",
            )
        bullets = exp.get("bullets_indices", [])
        if bullets:
            reemplazar_texto_parrafo(
                doc_plantilla.paragraphs[bullets[0]],
                f"{{{{EXP_{n}_BULLETS}}}}",
            )
            for idx in bullets[1:]:
                reemplazar_texto_parrafo(doc_plantilla.paragraphs[idx], "")

    salida_docx_plantilla.parent.mkdir(parents=True, exist_ok=True)
    doc_plantilla.save(str(salida_docx_plantilla))

    # 8. Guardar hash del DOCX original en metadatos (si la BD existe)
    if ruta_bd.exists():
        hash_docx = hashlib.sha256(cv_docx.read_bytes()).hexdigest()
        guardar_hash_cv(ruta_bd, hash_docx)

    # 9. Reportar
    print(f"DOCX pulido guardado en {salida_docx_pulido}")
    print(f"PDF pulido guardado en {salida_pdf}")
    print(f"Plantilla con placeholders guardada en {salida_docx_plantilla}")
    print(f"Correcciones aplicadas: {n_pulido}")
    print(f"Costo de la operación: ${costo:.4f} USD")


def _confirmar_interactivo(datos: dict) -> bool:
    print("\n=== Correcciones propuestas ===")
    for c in datos.get("correcciones", []):
        print(f"  - '{c['buscar']}' -> '{c['reemplazar']}'")
    indices_perfil = datos.get("bloque_perfil", [])
    print(f"\n=== Bloque Perfil identificado ===")
    print(f"  Párrafos {indices_perfil} (se reemplazarán con {{{{PERFIL}}}})")
    print(f"\n=== Experiencias identificadas ({len(datos.get('experiencias', []))}) ===")
    for n, exp in enumerate(datos.get("experiencias", []), start=1):
        print(f"  EXP_{n}: titulo idx {exp.get('titulo_idx')}, bullets {exp.get('bullets_indices')}")
    if datos.get("advertencias"):
        print("\n=== Advertencias ===")
        for adv in datos["advertencias"]:
            print(f"  [!] {adv}")
    respuesta = input("\n¿Aceptas estos cambios y guardas la plantilla? [s/N]: ").strip().lower()
    return respuesta == "s"


if __name__ == "__main__":
    ejecutar(confirmar=_confirmar_interactivo)
