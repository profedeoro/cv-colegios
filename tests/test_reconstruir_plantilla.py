from pathlib import Path
from unittest.mock import patch
from docx import Document
import pytest


def _crear_docx_fixture(tmp_path: Path) -> Path:
    """Crea un DOCX de prueba con secciones simuladas."""
    doc = Document()
    doc.add_paragraph("Daniel Eduardo Villalba")  # 0
    doc.add_paragraph("")                          # 1
    doc.add_paragraph("Perfil")                    # 2
    doc.add_paragraph("Soy docente con experiencia.")  # 3
    doc.add_paragraph("")                          # 4
    doc.add_paragraph("INTITUCIÓN EDUCATIVA X")    # 5 (typo intencional)
    doc.add_paragraph("Experiencia")               # 6
    doc.add_paragraph("Colegio ABC, 2024")         # 7 (titulo de experiencia)
    doc.add_paragraph("Enseñé matemáticas.")       # 8 (bullet)
    doc.add_paragraph("Coordiné actividades.")     # 9 (bullet)
    ruta = tmp_path / "cv.docx"
    doc.save(ruta)
    return ruta


def test_reconstruir_plantilla_genera_los_3_archivos(tmp_path):
    cv_docx = _crear_docx_fixture(tmp_path)
    salida_plantilla = tmp_path / "plantilla.docx"
    salida_pulido_docx = tmp_path / "pulido.docx"
    salida_pdf = tmp_path / "pulido.pdf"
    ruta_bd_inexistente = tmp_path / "no_existe.db"

    respuesta_simulada = (
        '{"correcciones": [{"buscar": "INTITUCIÓN", "reemplazar": "INSTITUCIÓN"}],'
        ' "bloque_perfil": [3],'
        ' "experiencias": [{"titulo_idx": 7, "bullets_indices": [8, 9]}],'
        ' "advertencias": []}'
    )

    with patch("reconstruir_plantilla.ClienteClaude") as mock_cls:
        cliente = mock_cls.return_value
        cliente.preguntar.return_value = (respuesta_simulada, 0.01)
        from reconstruir_plantilla import ejecutar
        ejecutar(
            cv_docx=cv_docx,
            salida_docx_plantilla=salida_plantilla,
            salida_docx_pulido=salida_pulido_docx,
            salida_pdf=salida_pdf,
            api_key="sk-test",
            confirmar=lambda _: True,
            ruta_bd=ruta_bd_inexistente,
        )

    assert salida_plantilla.exists()
    assert salida_pulido_docx.exists()
    assert salida_pdf.exists()


def test_reconstruir_plantilla_aplica_correcciones_en_pulido(tmp_path):
    cv_docx = _crear_docx_fixture(tmp_path)
    salida_plantilla = tmp_path / "plantilla.docx"
    salida_pulido_docx = tmp_path / "pulido.docx"
    salida_pdf = tmp_path / "pulido.pdf"

    respuesta_simulada = (
        '{"correcciones": [{"buscar": "INTITUCIÓN", "reemplazar": "INSTITUCIÓN"}],'
        ' "bloque_perfil": [3],'
        ' "experiencias": [{"titulo_idx": 7, "bullets_indices": [8, 9]}],'
        ' "advertencias": []}'
    )

    with patch("reconstruir_plantilla.ClienteClaude") as mock_cls:
        cliente = mock_cls.return_value
        cliente.preguntar.return_value = (respuesta_simulada, 0.01)
        from reconstruir_plantilla import ejecutar
        ejecutar(
            cv_docx=cv_docx,
            salida_docx_plantilla=salida_plantilla,
            salida_docx_pulido=salida_pulido_docx,
            salida_pdf=salida_pdf,
            api_key="sk-test",
            confirmar=lambda _: True,
            ruta_bd=tmp_path / "no.db",
        )

    doc_pulido = Document(salida_pulido_docx)
    textos_pulido = [p.text for p in doc_pulido.paragraphs]
    assert any("INSTITUCIÓN EDUCATIVA X" in t for t in textos_pulido)
    assert not any("INTITUCIÓN" in t for t in textos_pulido)


def test_reconstruir_plantilla_inserta_placeholders(tmp_path):
    cv_docx = _crear_docx_fixture(tmp_path)
    salida_plantilla = tmp_path / "plantilla.docx"
    salida_pulido_docx = tmp_path / "pulido.docx"
    salida_pdf = tmp_path / "pulido.pdf"

    respuesta_simulada = (
        '{"correcciones": [],'
        ' "bloque_perfil": [3],'
        ' "experiencias": [{"titulo_idx": 7, "bullets_indices": [8, 9]}],'
        ' "advertencias": []}'
    )

    with patch("reconstruir_plantilla.ClienteClaude") as mock_cls:
        cliente = mock_cls.return_value
        cliente.preguntar.return_value = (respuesta_simulada, 0.01)
        from reconstruir_plantilla import ejecutar
        ejecutar(
            cv_docx=cv_docx,
            salida_docx_plantilla=salida_plantilla,
            salida_docx_pulido=salida_pulido_docx,
            salida_pdf=salida_pdf,
            api_key="sk-test",
            confirmar=lambda _: True,
            ruta_bd=tmp_path / "no.db",
        )

    doc_plantilla = Document(salida_plantilla)
    textos = [p.text for p in doc_plantilla.paragraphs]
    assert "{{PERFIL}}" in textos
    assert "{{EXP_1_TITULO}}" in textos
    assert "{{EXP_1_BULLETS}}" in textos


def test_reconstruir_plantilla_acepta_respuesta_con_code_fence(tmp_path):
    """JSON envuelto en ```JSON...```` debe parsearse igual."""
    cv_docx = _crear_docx_fixture(tmp_path)
    salida_plantilla = tmp_path / "plantilla.docx"
    salida_pulido_docx = tmp_path / "pulido.docx"
    salida_pdf = tmp_path / "pulido.pdf"

    respuesta_con_fence = (
        '```JSON\n'
        '{"correcciones": [], "bloque_perfil": [3],'
        ' "experiencias": [{"titulo_idx": 7, "bullets_indices": [8]}],'
        ' "advertencias": []}'
        '\n```'
    )

    with patch("reconstruir_plantilla.ClienteClaude") as mock_cls:
        cliente = mock_cls.return_value
        cliente.preguntar.return_value = (respuesta_con_fence, 0.01)
        from reconstruir_plantilla import ejecutar
        ejecutar(
            cv_docx=cv_docx,
            salida_docx_plantilla=salida_plantilla,
            salida_docx_pulido=salida_pulido_docx,
            salida_pdf=salida_pdf,
            api_key="sk-test",
            confirmar=lambda _: True,
            ruta_bd=tmp_path / "no.db",
        )
    doc = Document(salida_plantilla)
    textos = [p.text for p in doc.paragraphs]
    assert "{{PERFIL}}" in textos


def test_falla_si_cv_no_existe(tmp_path):
    salida = tmp_path / "x.docx"
    pdf = tmp_path / "x.pdf"
    from reconstruir_plantilla import ejecutar
    with pytest.raises(FileNotFoundError, match="No se encontró"):
        ejecutar(
            cv_docx=tmp_path / "no_existe.docx",
            salida_docx_plantilla=salida,
            salida_docx_pulido=tmp_path / "pul.docx",
            salida_pdf=pdf,
            api_key="sk-test",
            confirmar=lambda _: True,
        )
