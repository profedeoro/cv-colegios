from pathlib import Path
import pytest
from docx import Document
from modulos.pdf_conversor import convertir_docx_a_pdf
from modulos.pdf_lector import leer_pdf


@pytest.fixture
def pdf_de_prueba(tmp_path):
    """Genera un PDF de prueba para no depender de un archivo binario en el repo."""
    docx_path = tmp_path / "test.docx"
    pdf_path = tmp_path / "test.pdf"
    doc = Document()
    doc.add_paragraph("Daniel Eduardo Villalba")
    doc.add_paragraph("Licenciado en educación física, 2024")
    doc.add_paragraph("ISBN: 978-99993-2-001-6")
    doc.save(docx_path)
    convertir_docx_a_pdf(docx_path, pdf_path)
    return pdf_path


def test_leer_pdf_extrae_texto(pdf_de_prueba):
    texto = leer_pdf(pdf_de_prueba)
    assert "Daniel Eduardo Villalba" in texto
    assert "2024" in texto
    assert "978-99993-2-001-6" in texto
