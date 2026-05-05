from pathlib import Path
import pytest
from docx import Document
from modulos.docx_editor import (
    enumerar_parrafos,
    aplicar_correcciones_a_parrafo,
    aplicar_correcciones_a_documento,
    reemplazar_texto_parrafo,
)


def _docx_con_estilos(tmp_path) -> Path:
    """Crea un DOCX con varios párrafos y estilos para tests."""
    doc = Document()
    doc.add_paragraph("")  # vacío
    doc.add_paragraph("Daniel Eduardo Villalba")
    doc.add_paragraph("INTITUCIÓN UNIVERSITARIA POLITECNICO")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Soy docente con experiencia en ")
    r2 = p.add_run("phyton")
    r2.bold = True
    p.add_run(" para análisis estadístico.")
    ruta = tmp_path / "doc.docx"
    doc.save(ruta)
    return ruta


def test_enumerar_parrafos_omite_vacios(tmp_path):
    ruta = _docx_con_estilos(tmp_path)
    doc = Document(ruta)
    pares = enumerar_parrafos(doc)
    indices = [i for i, _ in pares]
    textos = [t for _, t in pares]
    assert 0 not in indices  # vacío
    assert 1 in indices and "Daniel" in textos[indices.index(1)]
    assert 2 in indices and "INTITUCIÓN" in textos[indices.index(2)]
    assert 3 not in indices  # vacío
    assert 4 in indices


def test_aplicar_correcciones_a_parrafo_corrige_typo_simple(tmp_path):
    ruta = _docx_con_estilos(tmp_path)
    doc = Document(ruta)
    parrafo = doc.paragraphs[2]  # "INTITUCIÓN UNIVERSITARIA POLITECNICO"
    n = aplicar_correcciones_a_parrafo(parrafo, [("INTITUCIÓN", "INSTITUCIÓN")])
    assert n == 1
    assert "INSTITUCIÓN UNIVERSITARIA POLITECNICO" == parrafo.text


def test_aplicar_correcciones_preserva_estilo_de_runs(tmp_path):
    ruta = _docx_con_estilos(tmp_path)
    doc = Document(ruta)
    parrafo = doc.paragraphs[4]  # "Soy docente con experiencia en phyton para análisis..."

    # phyton está en run[1] que es bold. Aplicar la corrección:
    n = aplicar_correcciones_a_parrafo(parrafo, [("phyton", "Python")])
    assert n == 1
    assert "Python" in parrafo.text
    # Verificar que el run que tenía phyton ahora dice Python y SIGUE siendo bold:
    runs_bold = [r for r in parrafo.runs if r.bold]
    assert any("Python" in r.text for r in runs_bold)


def test_aplicar_correcciones_a_documento_recorre_todos(tmp_path):
    ruta = _docx_con_estilos(tmp_path)
    doc = Document(ruta)
    n = aplicar_correcciones_a_documento(doc, [
        ("INTITUCIÓN", "INSTITUCIÓN"),
        ("phyton", "Python"),
    ])
    assert n == 2  # uno por cada corrección
    textos = [p.text for p in doc.paragraphs]
    assert any("INSTITUCIÓN" in t for t in textos)
    assert any("Python" in t for t in textos)
    assert not any("INTITUCIÓN" in t for t in textos)
    assert not any("phyton" in t for t in textos)


def test_aplicar_correcciones_recorre_tablas(tmp_path):
    """Las correcciones también deben aplicarse a celdas de tablas."""
    doc = Document()
    tabla = doc.add_table(rows=1, cols=2)
    tabla.cell(0, 0).paragraphs[0].add_run("INTITUCIÓN X")
    tabla.cell(0, 1).paragraphs[0].add_run("Algo más")
    ruta = tmp_path / "tabla.docx"
    doc.save(ruta)

    doc2 = Document(ruta)
    n = aplicar_correcciones_a_documento(doc2, [("INTITUCIÓN", "INSTITUCIÓN")])
    assert n == 1
    assert "INSTITUCIÓN" in doc2.tables[0].cell(0, 0).text


def test_reemplazar_texto_parrafo_pone_todo_en_run0(tmp_path):
    ruta = _docx_con_estilos(tmp_path)
    doc = Document(ruta)
    parrafo = doc.paragraphs[4]  # tenía 3 runs
    reemplazar_texto_parrafo(parrafo, "{{PERFIL}}")
    assert parrafo.text == "{{PERFIL}}"
    # El primer run ahora tiene el placeholder, los demás están vacíos:
    assert parrafo.runs[0].text == "{{PERFIL}}"
    for r in parrafo.runs[1:]:
        assert r.text == ""


def test_reemplazar_texto_parrafo_preserva_estilo_de_run0(tmp_path):
    """Si el run[0] tenía bold=True, el placeholder también queda bold."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Original bold")
    r.bold = True
    p.add_run(" no bold")
    ruta = tmp_path / "neg.docx"
    doc.save(ruta)

    doc2 = Document(ruta)
    parrafo = doc2.paragraphs[0]
    reemplazar_texto_parrafo(parrafo, "{{NEW}}")
    assert parrafo.runs[0].text == "{{NEW}}"
    assert parrafo.runs[0].bold is True


def test_reemplazar_en_parrafo_sin_runs_agrega_uno(tmp_path):
    doc = Document()
    p = doc.add_paragraph()  # sin runs
    reemplazar_texto_parrafo(p, "{{X}}")
    assert p.text == "{{X}}"
