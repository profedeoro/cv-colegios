from pathlib import Path
import pytest
from docx import Document
from modulos.plantilla import rellenar_plantilla


def _crear_plantilla(tmp_path) -> Path:
    doc = Document()
    doc.add_heading("{{NOMBRE}}", level=1)
    doc.add_paragraph("{{PERFIL}}")
    doc.add_heading("Experiencia", level=2)
    doc.add_paragraph("{{EXP_1_TITULO}}")
    p = tmp_path / "plantilla.docx"
    doc.save(p)
    return p


def test_rellenar_plantilla_reemplaza_placeholders(tmp_path):
    plantilla = _crear_plantilla(tmp_path)
    salida = tmp_path / "salida.docx"

    rellenar_plantilla(
        plantilla,
        salida,
        valores={
            "NOMBRE": "Daniel Villalba",
            "PERFIL": "Docente con experiencia en TIC.",
            "EXP_1_TITULO": "Colegio Inocencio Chincá",
        },
    )

    doc = Document(salida)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Daniel Villalba" in texto
    assert "Docente con experiencia en TIC." in texto
    assert "Colegio Inocencio Chincá" in texto
    assert "{{NOMBRE}}" not in texto


def test_rellenar_falla_si_falta_valor(tmp_path):
    plantilla = _crear_plantilla(tmp_path)
    salida = tmp_path / "salida.docx"
    with pytest.raises(ValueError, match="placeholder"):
        rellenar_plantilla(plantilla, salida, valores={"NOMBRE": "X"})


def test_rellenar_distribuye_multilinea_en_parrafos_vacios(tmp_path):
    """Multi-line value debe repartirse en placeholder + párrafos vacíos siguientes."""
    from docx import Document
    plantilla = tmp_path / "p.docx"
    salida = tmp_path / "s.docx"
    doc = Document()
    doc.add_paragraph("{{BULLETS}}")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("FIN")
    doc.save(str(plantilla))

    rellenar_plantilla(plantilla, salida, {"BULLETS": "linea1\nlinea2\nlinea3"})

    out = Document(str(salida))
    textos = [p.text for p in out.paragraphs]
    assert textos[0] == "linea1"
    assert textos[1] == "linea2"
    assert textos[2] == "linea3"
    assert textos[3] == "FIN"


def test_rellenar_agrupa_excedente_si_faltan_slots(tmp_path):
    """Más líneas que párrafos vacíos: las sobrantes se concatenan en el último slot."""
    from docx import Document
    plantilla = tmp_path / "p.docx"
    salida = tmp_path / "s.docx"
    doc = Document()
    doc.add_paragraph("{{BULLETS}}")
    doc.add_paragraph("")  # solo 1 slot vacío
    doc.add_paragraph("siguiente")
    doc.save(str(plantilla))

    rellenar_plantilla(plantilla, salida, {"BULLETS": "a\nb\nc\nd"})

    out = Document(str(salida))
    textos = [p.text for p in out.paragraphs]
    assert textos[0] == "a"
    # 'b', 'c', 'd' deben acabar agrupadas en el slot 1 (que era el último vacío)
    assert textos[1] == "b\nc\nd"
    assert textos[2] == "siguiente"


def test_rellenar_no_distribuye_si_valor_es_una_linea(tmp_path):
    """Valor de una línea NO debe tocar párrafos siguientes."""
    from docx import Document
    plantilla = tmp_path / "p.docx"
    salida = tmp_path / "s.docx"
    doc = Document()
    doc.add_paragraph("{{X}}")
    doc.add_paragraph("")
    doc.add_paragraph("Y")
    doc.save(str(plantilla))

    rellenar_plantilla(plantilla, salida, {"X": "valor_simple"})

    out = Document(str(salida))
    assert out.paragraphs[0].text == "valor_simple"
    assert out.paragraphs[1].text == ""  # NO se llena
    assert out.paragraphs[2].text == "Y"


def test_rellenar_omite_parrafos_consumidos(tmp_path):
    """Después de distribuir, los párrafos vacíos consumidos no deben ser revisados de nuevo.

    Importante: si después de los empties hay OTRO placeholder, ese placeholder
    debe procesarse normalmente.
    """
    from docx import Document
    plantilla = tmp_path / "p.docx"
    salida = tmp_path / "s.docx"
    doc = Document()
    doc.add_paragraph("{{A}}")
    doc.add_paragraph("")
    doc.add_paragraph("{{B}}")
    doc.save(str(plantilla))

    rellenar_plantilla(plantilla, salida, {"A": "a1\na2", "B": "b_solo"})

    out = Document(str(salida))
    textos = [p.text for p in out.paragraphs]
    assert textos[0] == "a1"
    assert textos[1] == "a2"
    assert textos[2] == "b_solo"
