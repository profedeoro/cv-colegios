"""Tests del orquestador `modulos/generar.py`.

Cubre:
- _extraer_cv_base: extracción del PERFIL + EXP_N_TITULO/BULLETS desde el polished docx
  usando la plantilla como guía.
- _intentar_con_validacion: bucle de hasta 3 intentos con detector de alucinaciones.
- procesar_colegio: happy path, retry tras alucinación, giveup tras 3 intentos.
- ejecutar: procesa varios colegios, registra la ejecución, acumula costos.
"""
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from modulos.db import (
    inicializar_db, insertar_colegio, marcar_enriquecido,
    conectar, colegios_para_generar,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

PERFIL_PEDAGOGICO_EJEMPLO = {
    "bilingue": False,
    "idioma_segundo": None,
    "religioso": False,
    "denominacion": None,
    "ib": False,
    "montessori": False,
    "enfoque_deportivo": True,
    "enfoque_tecnico": False,
    "enfasis_tic": True,
    "tamano_estimado": "mediano",
    "palabras_clave": ["innovación", "deporte"],
}


def _crear_plantilla_mini(path: Path) -> None:
    """Plantilla pequeña con PERFIL + EXP_1_TITULO + EXP_1_BULLETS."""
    doc = Document()
    doc.add_paragraph("Daniel Eduardo Villalba de Oro")
    doc.add_paragraph("Perfil")
    doc.add_paragraph("{{PERFIL}}")
    doc.add_paragraph("EXPERIENCIA")
    doc.add_paragraph("{{EXP_1_TITULO}}")
    doc.add_paragraph("{{EXP_1_BULLETS}}")
    doc.add_paragraph("")
    doc.add_paragraph("FIN")
    doc.save(str(path))


def _crear_polished_mini(path: Path) -> None:
    """Polished CV con la misma estructura, pero con texto real en lugar de placeholders."""
    doc = Document()
    doc.add_paragraph("Daniel Eduardo Villalba de Oro")
    doc.add_paragraph("Perfil")
    doc.add_paragraph(
        "Profesional con experiencia en docencia y TICs. Investigador "
        "en innovación educativa con publicaciones indexadas."
    )
    doc.add_paragraph("EXPERIENCIA")
    doc.add_paragraph("INSTITUCIÓN EDUCATIVA INOCENCIO CHINCÁ — Barranquilla")
    doc.add_paragraph("Diseñé planes de educación física.")
    doc.add_paragraph("Implementé evaluación formativa.")
    doc.add_paragraph("FIN")
    doc.save(str(path))


def _bd_con_colegio_enriquecido(tmp_path: Path, *, nombre: str = "Colegio Test",
                                ciudad: str = "Bogotá") -> tuple[Path, int]:
    bd = tmp_path / "t.db"
    inicializar_db(bd)
    cid = insertar_colegio(
        bd, nombre=nombre, ciudad=ciudad,
        departamento="Bogotá D.C.", fuente="MEN",
    )
    marcar_enriquecido(
        bd, cid,
        web="https://x.edu.co",
        correo="info@x.edu.co",
        correo_destinatario="rector@x.edu.co",
        perfil_pedagogico=PERFIL_PEDAGOGICO_EJEMPLO,
        palabras_clave=PERFIL_PEDAGOGICO_EJEMPLO["palabras_clave"],
    )
    return bd, cid


def _colegio_dict(bd: Path, cid: int) -> dict:
    """Carga el colegio como dict (tal como lo entrega colegios_para_generar)."""
    pendientes = colegios_para_generar(bd, limite=10)
    return next(c for c in pendientes if c["id"] == cid)


# ---------------------------------------------------------------------------
# _extraer_cv_base
# ---------------------------------------------------------------------------

def test_extraer_cv_base_mapea_placeholder_a_polished(tmp_path):
    from modulos.generar import _extraer_cv_base
    plantilla = tmp_path / "plantilla.docx"
    polished = tmp_path / "polished.docx"
    _crear_plantilla_mini(plantilla)
    _crear_polished_mini(polished)

    valores = _extraer_cv_base(plantilla, polished)

    assert "PERFIL" in valores
    assert "innovación educativa" in valores["PERFIL"]
    assert valores["EXP_1_TITULO"] == "INSTITUCIÓN EDUCATIVA INOCENCIO CHINCÁ — Barranquilla"
    bullets = valores["EXP_1_BULLETS"]
    assert "Diseñé planes" in bullets
    assert "Implementé evaluación" in bullets
    # bullets se entrega como string multilínea
    assert bullets.count("\n") == 1


def test_extraer_cv_base_extrae_todos_los_placeholders_de_la_plantilla_real():
    """La plantilla real (data/plantilla_base.docx) tiene 13 placeholders. Todos deben extraerse."""
    from modulos.generar import _extraer_cv_base
    raiz = Path(__file__).parent.parent
    plantilla = raiz / "data" / "plantilla_base.docx"
    polished = raiz / "data" / "cv_base_polished.docx"
    if not plantilla.exists() or not polished.exists():
        pytest.skip("Plantilla / polished no presentes en este entorno")
    valores = _extraer_cv_base(plantilla, polished)
    esperados = {"PERFIL"} | {f"EXP_{n}_{kind}" for n in range(1, 7) for kind in ("TITULO", "BULLETS")}
    assert set(valores.keys()) == esperados
    # Todos no vacíos
    for k, v in valores.items():
        assert v.strip(), f"{k} salió vacío"


# ---------------------------------------------------------------------------
# _intentar_con_validacion
# ---------------------------------------------------------------------------

def test_intentar_con_validacion_exito_primer_intento():
    from modulos.generar import _intentar_con_validacion

    def generador():
        return ("Daniel Eduardo Villalba de Oro es docente.", 0.01)

    cv_original = "Daniel Eduardo Villalba de Oro es licenciado en educación física."
    texto, costo, intentos = _intentar_con_validacion(
        generador, cv_original=cv_original, nombres_permitidos=set(), max_intentos=3,
    )
    assert "Daniel" in texto
    assert costo == pytest.approx(0.01)
    assert intentos == 1


def test_intentar_con_validacion_recupera_tras_alucinacion():
    from modulos.generar import _intentar_con_validacion

    respuestas = iter([
        ("Trabajó en Microsoft Research en 2001.", 0.01),  # alucina (Microsoft Research, 2001)
        ("Daniel es docente con experiencia en TICs.", 0.02),  # ok
    ])

    def generador():
        return next(respuestas)

    cv_original = (
        "Daniel Eduardo Villalba de Oro es docente con experiencia en TICs. "
        "Tiene un libro con ISBN 978-99993-2-001-6."
    )
    texto, costo, intentos = _intentar_con_validacion(
        generador, cv_original=cv_original, nombres_permitidos=set(), max_intentos=3,
    )
    assert "Daniel" in texto
    assert "Microsoft" not in texto
    assert costo == pytest.approx(0.03)
    assert intentos == 2


def test_intentar_con_validacion_giveup_tras_3_intentos():
    from modulos.generar import _intentar_con_validacion

    def generador():
        return ("Inventó la rueda en 1492 junto a Colón.", 0.01)

    cv_original = "Daniel es docente."
    resultado = _intentar_con_validacion(
        generador, cv_original=cv_original, nombres_permitidos=set(), max_intentos=3,
    )
    assert resultado is None  # señal de giveup


def test_nombres_permitidos_incluye_nombre_completo_y_ciudad():
    """_nombres_permitidos debe incluir las frases de nombre propio extraídas
    (p. ej. "San José", "Bogotá") y, como fallback, el nombre completo y la
    ciudad completa, para que el validador acepte frases multi-palabra."""
    from modulos.generar import _nombres_permitidos

    colegio = {"nombre": "Colegio Bilingüe San José", "ciudad": "Bogotá"}
    permitidos = _nombres_permitidos(colegio)
    # Frases de nombre propio extraídas por `extraer_hechos`
    assert "San José" in permitidos
    assert "Bogotá" in permitidos
    # Fallback: el nombre completo del colegio también debe estar
    assert "Colegio Bilingüe San José" in permitidos


def test_intentar_con_validacion_acepta_nombres_del_colegio_en_la_carta():
    """Cuando la carta menciona el nombre del colegio + ciudad, el validador no debe
    rechazarla aun si esos tokens no aparecen en el CV de Daniel."""
    from modulos.generar import _intentar_con_validacion, _nombres_permitidos

    def generador():
        return ("Postulación al Colegio Bilingüe San José en Bogotá.", 0.01)

    cv_original = "Daniel es docente."
    permitidos = _nombres_permitidos({"nombre": "Colegio Bilingüe San José", "ciudad": "Bogotá"})
    resultado = _intentar_con_validacion(
        generador,
        cv_original=cv_original,
        nombres_permitidos=permitidos,
        max_intentos=3,
    )
    assert resultado is not None, "El validador no debe rechazar el nombre del colegio"
    texto, _, _ = resultado
    assert "Colegio" in texto


# ---------------------------------------------------------------------------
# procesar_colegio - happy path
# ---------------------------------------------------------------------------

def _cliente_claude_secuencial(respuestas: list[tuple[str, float]]):
    """Mock de cliente_claude.preguntar que devuelve respuestas en orden."""
    it = iter(respuestas)
    mock = MagicMock()
    mock.preguntar.side_effect = lambda **kw: next(it)
    return mock


def test_procesar_colegio_happy_path(tmp_path, monkeypatch):
    """Caso feliz: 3 llamadas a Claude OK al primer intento, PDF mockeado, borrador insertado."""
    from modulos import generar

    bd, cid = _bd_con_colegio_enriquecido(tmp_path, nombre="Colegio Las Flores")
    colegio = _colegio_dict(bd, cid)

    plantilla = tmp_path / "plantilla.docx"
    polished_docx = tmp_path / "polished.docx"
    _crear_plantilla_mini(plantilla)
    _crear_polished_mini(polished_docx)
    polished_pdf = tmp_path / "polished.pdf"
    polished_pdf.write_text("dummy", encoding="utf-8")

    # Cualquier texto vale; el validador usa el CV original. Devolvemos texto
    # que sólo contiene palabras presentes en el CV original (no aluce).
    cliente = _cliente_claude_secuencial([
        ("Docente con experiencia en TICs e innovación educativa.", 0.01),  # perfil
        ("Diseñé planes de educación física.\nImplementé evaluación formativa.", 0.02),  # bullets
        ("Estimado equipo directivo, soy Daniel Villalba, docente de educación física.", 0.03),  # carta
    ])

    # Mock del texto del CV (para no llamar a pdfplumber)
    monkeypatch.setattr(
        generar, "_texto_cv_completo",
        lambda _p: (
            "Daniel Eduardo Villalba de Oro Docente con experiencia en TICs e innovación educativa. "
            "Diseñé planes de educación física. Implementé evaluación formativa. "
            "Estimado equipo directivo soy Daniel Villalba docente de educación física."
        ),
    )

    # Mock de la conversión a PDF para no invocar LibreOffice
    pdfs_generados = []
    def fake_pdf(plantilla_path, valores, salida_pdf):
        Path(salida_pdf).parent.mkdir(parents=True, exist_ok=True)
        Path(salida_pdf).write_bytes(b"%PDF-1.4 fake\n")
        pdfs_generados.append((Path(salida_pdf), dict(valores)))
    monkeypatch.setattr(generar, "generar_pdf_personalizado", fake_pdf)

    salida_dir = tmp_path / "salida"
    resultado = generar.procesar_colegio(
        bd, colegio, cliente,
        plantilla_path=plantilla,
        polished_pdf_path=polished_pdf,
        polished_docx_path=polished_docx,
        salida_dir=salida_dir,
    )

    assert resultado["estado_final"] == "borrador_insertado"
    assert resultado["costo"] == pytest.approx(0.01 + 0.02 + 0.03, abs=1e-6)
    # PDF creado en ruta esperada
    assert len(pdfs_generados) == 1
    ruta_pdf, valores = pdfs_generados[0]
    assert ruta_pdf.exists()
    assert ruta_pdf.name.endswith(".pdf")
    # El slug viene de nombre_normalizado del colegio
    assert "flores" in ruta_pdf.name.lower()
    # Valores deben contener todos los placeholders de la plantilla
    assert "PERFIL" in valores
    assert "EXP_1_TITULO" in valores
    assert "EXP_1_BULLETS" in valores

    # Borrador insertado
    conn = conectar(bd)
    try:
        row = conn.execute(
            "SELECT colegio_id, tipo, asunto, cuerpo_carta, ruta_pdf_hv, estado FROM borradores"
        ).fetchone()
    finally:
        conn.close()
    assert row is not None
    assert row["colegio_id"] == cid
    assert row["tipo"] == "inicial"
    assert "Colegio Las Flores" in row["asunto"]
    assert "Daniel Villalba" in row["cuerpo_carta"]
    assert row["ruta_pdf_hv"] == str(ruta_pdf)
    assert row["estado"] == "listo_para_subir"

    # Colegio sigue en 'enriquecido' (transición a borrador_creado pasa en Tarea 7)
    conn = conectar(bd)
    try:
        estado_actual = conn.execute("SELECT estado FROM colegios WHERE id = ?", (cid,)).fetchone()[0]
    finally:
        conn.close()
    assert estado_actual == "enriquecido"


def test_procesar_colegio_reintenta_si_aluce(tmp_path, monkeypatch):
    """Si Claude alucina 2 veces y luego responde bien, el borrador se inserta igual."""
    from modulos import generar

    bd, cid = _bd_con_colegio_enriquecido(tmp_path)
    colegio = _colegio_dict(bd, cid)

    plantilla = tmp_path / "plantilla.docx"
    polished_docx = tmp_path / "polished.docx"
    _crear_plantilla_mini(plantilla)
    _crear_polished_mini(polished_docx)
    polished_pdf = tmp_path / "polished.pdf"
    polished_pdf.write_text("dummy", encoding="utf-8")

    cliente = _cliente_claude_secuencial([
        # Perfil: alucina (Microsoft Research) y luego ok
        ("Trabajó en Microsoft Research en 2001.", 0.01),
        ("Docente con experiencia en TICs e innovación educativa.", 0.01),
        # Bullets: ok primer intento (texto sin números/nombres nuevos)
        ("Diseñé planes de educación física.\nImplementé evaluación formativa.", 0.02),
        # Carta: ok primer intento
        ("Estimado equipo directivo soy Daniel Villalba docente.", 0.03),
    ])

    monkeypatch.setattr(
        generar, "_texto_cv_completo",
        lambda _p: (
            "Daniel Eduardo Villalba de Oro docente con experiencia en TICs e innovación educativa. "
            "Diseñé planes de educación física. Implementé evaluación formativa. "
            "Estimado equipo directivo soy Daniel Villalba docente."
        ),
    )

    pdfs = []
    monkeypatch.setattr(
        generar, "generar_pdf_personalizado",
        lambda p, v, s: (Path(s).parent.mkdir(parents=True, exist_ok=True),
                         Path(s).write_bytes(b"%PDF-1.4"),
                         pdfs.append(Path(s)))[0],
    )

    salida_dir = tmp_path / "salida"
    resultado = generar.procesar_colegio(
        bd, colegio, cliente,
        plantilla_path=plantilla,
        polished_pdf_path=polished_pdf,
        polished_docx_path=polished_docx,
        salida_dir=salida_dir,
    )

    assert resultado["estado_final"] == "borrador_insertado"
    # Claude llamado 4 veces (1 perfil-retry + 1 perfil-ok + 1 bullets + 1 carta)
    assert cliente.preguntar.call_count == 4

    # Borrador insertado
    conn = conectar(bd)
    try:
        n_borradores = conn.execute("SELECT COUNT(*) FROM borradores WHERE colegio_id = ?", (cid,)).fetchone()[0]
    finally:
        conn.close()
    assert n_borradores == 1


def test_procesar_colegio_giveup_marca_revisar_manualmente(tmp_path, monkeypatch):
    """Si Claude alucina 3 veces seguidas en cualquier paso, el colegio va a revisar_manualmente."""
    from modulos import generar

    bd, cid = _bd_con_colegio_enriquecido(tmp_path)
    colegio = _colegio_dict(bd, cid)

    plantilla = tmp_path / "plantilla.docx"
    polished_docx = tmp_path / "polished.docx"
    _crear_plantilla_mini(plantilla)
    _crear_polished_mini(polished_docx)
    polished_pdf = tmp_path / "polished.pdf"
    polished_pdf.write_text("dummy", encoding="utf-8")

    # Perfil: alucina las 3 veces.
    cliente = _cliente_claude_secuencial([
        ("Inventó la rueda en 1492.", 0.01),
        ("Trabajó con Albert Einstein en 1922.", 0.01),
        ("Lideró el Proyecto Manhattan en 1943.", 0.01),
    ])

    monkeypatch.setattr(
        generar, "_texto_cv_completo",
        lambda _p: "Daniel Eduardo Villalba de Oro docente.",
    )

    pdfs = []
    monkeypatch.setattr(
        generar, "generar_pdf_personalizado",
        lambda p, v, s: pdfs.append(Path(s)),
    )

    salida_dir = tmp_path / "salida"
    resultado = generar.procesar_colegio(
        bd, colegio, cliente,
        plantilla_path=plantilla,
        polished_pdf_path=polished_pdf,
        polished_docx_path=polished_docx,
        salida_dir=salida_dir,
    )

    assert resultado["estado_final"] == "revisar_manualmente"
    # No se generó PDF
    assert pdfs == []
    # No hay borrador
    conn = conectar(bd)
    try:
        n_borradores = conn.execute("SELECT COUNT(*) FROM borradores WHERE colegio_id = ?", (cid,)).fetchone()[0]
        row = conn.execute("SELECT estado, intentos_generar FROM colegios WHERE id = ?", (cid,)).fetchone()
    finally:
        conn.close()
    assert n_borradores == 0
    assert row["estado"] == "revisar_manualmente"
    # intentos_generar incrementado una vez (un colegio fallado, un intento de orquestación)
    assert row["intentos_generar"] == 1


# ---------------------------------------------------------------------------
# ejecutar
# ---------------------------------------------------------------------------

def test_ejecutar_procesa_varios_colegios(tmp_path, monkeypatch):
    """ejecutar() procesa N colegios secuencialmente y registra la ejecución."""
    from modulos import generar

    bd = tmp_path / "t.db"
    inicializar_db(bd)
    ids = []
    for i in range(2):
        cid = insertar_colegio(
            bd, nombre=f"Colegio Test {i}", ciudad="Bogotá",
            departamento="Bogotá D.C.", fuente="MEN",
        )
        marcar_enriquecido(
            bd, cid, web="https://x.edu.co",
            correo="info@x.edu.co", correo_destinatario="rector@x.edu.co",
            perfil_pedagogico=PERFIL_PEDAGOGICO_EJEMPLO,
            palabras_clave=PERFIL_PEDAGOGICO_EJEMPLO["palabras_clave"],
        )
        ids.append(cid)

    plantilla = tmp_path / "plantilla.docx"
    polished_docx = tmp_path / "polished.docx"
    _crear_plantilla_mini(plantilla)
    _crear_polished_mini(polished_docx)
    polished_pdf = tmp_path / "polished.pdf"
    polished_pdf.write_text("dummy", encoding="utf-8")

    # 3 llamadas por colegio × 2 colegios = 6 llamadas. Todas exitosas.
    respuestas = []
    for _ in range(2):
        respuestas.extend([
            ("Docente con experiencia en TICs e innovación educativa.", 0.01),  # perfil
            ("Diseñé planes de educación física.\nImplementé evaluación formativa.", 0.02),
            ("Estimado equipo directivo soy Daniel Villalba docente.", 0.03),
        ])
    cliente = _cliente_claude_secuencial(respuestas)

    monkeypatch.setattr(
        generar, "_texto_cv_completo",
        lambda _p: (
            "Daniel Eduardo Villalba de Oro docente con experiencia en TICs e innovación educativa. "
            "Diseñé planes de educación física. Implementé evaluación formativa. "
            "Estimado equipo directivo soy Daniel Villalba docente."
        ),
    )

    def fake_pdf(p, v, s):
        Path(s).parent.mkdir(parents=True, exist_ok=True)
        Path(s).write_bytes(b"%PDF-1.4")
    monkeypatch.setattr(generar, "generar_pdf_personalizado", fake_pdf)

    salida_dir = tmp_path / "salida"
    resumen = generar.ejecutar(
        bd, cliente, max_colegios=15,
        plantilla_path=plantilla,
        polished_pdf_path=polished_pdf,
        polished_docx_path=polished_docx,
        salida_dir=salida_dir,
    )

    assert resumen["resumen"]["borrador_insertado"] == 2
    assert resumen["costo_usd"] == pytest.approx(0.06 * 2, abs=1e-6)

    conn = conectar(bd)
    try:
        # 2 borradores creados
        n_b = conn.execute("SELECT COUNT(*) FROM borradores").fetchone()[0]
        # 1 ejecución registrada para módulo 'generar'
        n_ej = conn.execute(
            "SELECT COUNT(*) FROM registro_ejecuciones WHERE modulo = 'generar'"
        ).fetchone()[0]
        ej = conn.execute(
            "SELECT estado, colegios_procesados, costo_api_usd "
            "FROM registro_ejecuciones WHERE modulo = 'generar' "
            "ORDER BY id DESC LIMIT 1"
        ).fetchone()
    finally:
        conn.close()
    assert n_b == 2
    assert n_ej == 1
    assert ej["estado"] == "ok"
    assert ej["colegios_procesados"] == 2
    assert ej["costo_api_usd"] == pytest.approx(0.06 * 2, abs=1e-6)


def test_ejecutar_sin_pendientes_no_falla(tmp_path):
    """Si no hay colegios enriquecidos, ejecutar() no rompe y registra ejecución vacía."""
    from modulos import generar

    bd = tmp_path / "t.db"
    inicializar_db(bd)
    cliente = MagicMock()

    raiz = Path(__file__).parent.parent
    resumen = generar.ejecutar(
        bd, cliente, max_colegios=15,
        plantilla_path=raiz / "data" / "plantilla_base.docx",
        polished_pdf_path=raiz / "data" / "cv_base_polished.pdf",
        polished_docx_path=raiz / "data" / "cv_base_polished.docx",
        salida_dir=tmp_path / "salida",
    )
    assert resumen["resumen"].get("borrador_insertado", 0) == 0
    cliente.preguntar.assert_not_called()


# ---------------------------------------------------------------------------
# asunto del correo
# ---------------------------------------------------------------------------

def test_asunto_sigue_plantilla_del_spec(tmp_path, monkeypatch):
    """El asunto debe seguir el formato del spec: '... — {Nombre Colegio}'."""
    from modulos import generar

    bd, cid = _bd_con_colegio_enriquecido(tmp_path, nombre="Colegio Pedro Pascasio Martínez")
    colegio = _colegio_dict(bd, cid)

    plantilla = tmp_path / "plantilla.docx"
    polished_docx = tmp_path / "polished.docx"
    _crear_plantilla_mini(plantilla)
    _crear_polished_mini(polished_docx)
    polished_pdf = tmp_path / "polished.pdf"
    polished_pdf.write_text("dummy", encoding="utf-8")

    cliente = _cliente_claude_secuencial([
        ("Docente con experiencia en TICs e innovación educativa.", 0.01),
        ("Diseñé planes de educación física.\nImplementé evaluación formativa.", 0.02),
        ("Estimado equipo directivo soy Daniel Villalba docente.", 0.03),
    ])
    monkeypatch.setattr(
        generar, "_texto_cv_completo",
        lambda _p: (
            "Daniel Eduardo Villalba de Oro docente con experiencia en TICs e innovación educativa. "
            "Diseñé planes de educación física. Implementé evaluación formativa. "
            "Estimado equipo directivo soy Daniel Villalba docente."
        ),
    )
    monkeypatch.setattr(generar, "generar_pdf_personalizado", lambda p, v, s: Path(s).write_bytes(b"%PDF"))

    generar.procesar_colegio(
        bd, colegio, cliente,
        plantilla_path=plantilla,
        polished_pdf_path=polished_pdf,
        polished_docx_path=polished_docx,
        salida_dir=tmp_path / "salida",
    )

    conn = conectar(bd)
    try:
        asunto = conn.execute("SELECT asunto FROM borradores WHERE colegio_id = ?", (cid,)).fetchone()[0]
    finally:
        conn.close()
    # Sigue el formato del spec 4.3
    assert "Daniel" in asunto
    assert "Colegio Pedro Pascasio Martínez" in asunto


# ---------------------------------------------------------------------------
# _normalizar_nombre_colegio / _nombres_permitidos con nombres MEN abreviados
# ---------------------------------------------------------------------------

def test_normalizar_nombre_expande_abreviaturas_men():
    from modulos.generar import _normalizar_nombre_colegio
    assert _normalizar_nombre_colegio("COL FUND SANTA MARIA") == "Colegio Fundación Santa Maria"


def test_normalizar_nombre_titlecasea_all_caps_sin_abreviatura():
    from modulos.generar import _normalizar_nombre_colegio
    assert _normalizar_nombre_colegio("ESCUELA PEDAGOGICA INTEGRAL") == "Escuela Pedagogica Integral"


def test_normalizar_nombre_preserva_ya_normalizado():
    from modulos.generar import _normalizar_nombre_colegio
    assert _normalizar_nombre_colegio("Colegio Bilingüe San José") == "Colegio Bilingüe San José"


def test_nombres_permitidos_incluye_version_normalizada():
    """permitidos debe cubrir tanto el nombre crudo como el expandido."""
    from modulos.generar import _nombres_permitidos
    colegio = {"id": 1, "nombre": "COL FUND SANTA MARIA", "ciudad": "Bogotá,"}
    permitidos = _nombres_permitidos(colegio)
    # El nombre normalizado completo como string fallback
    assert "Colegio Fundación Santa Maria" in permitidos
    # Y la ciudad sin coma
    assert "Bogotá" in permitidos
    # El crudo también, por si Claude lo usa textualmente
    assert "COL FUND SANTA MARIA" in permitidos
