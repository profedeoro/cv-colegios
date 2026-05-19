"""Inspecciona la estructura de plantilla_base.docx: alineacion, bullets, runs."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("data/plantilla_base.docx")

alineaciones = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    None: "default",
}

# Buscar paragrafos con placeholders + sus adyacentes
print("=== Paragrafos con placeholders y sus 3 vecinos posteriores ===\n")
for i, p in enumerate(doc.paragraphs):
    if "{{" in p.text:
        print(f"[{i}] PLACEHOLDER: {p.text!r}")
        print(f"     align: {alineaciones.get(p.alignment, p.alignment)}")
        print(f"     style: {p.style.name}")
        # XML para ver direccion text
        xml = p._p.xml
        if "bidi" in xml or "rtl" in xml.lower():
            print(f"     [!] tiene marca RTL/bidi en XML")
        for k in range(1, 6):
            if i + k < len(doc.paragraphs):
                pp = doc.paragraphs[i + k]
                texto = pp.text.strip()
                marker = "EMPTY" if not texto else f"'{texto[:50]}'"
                ali = alineaciones.get(pp.alignment, pp.alignment)
                print(f"  [{i+k}] {marker} (align={ali}, style={pp.style.name})")
        print()

print("\n=== TOTAL paragrafos:", len(doc.paragraphs))
print("\n=== Estilos disponibles (relacionados a list/bullet) ===")
for s in doc.styles:
    nombre = s.name.lower()
    if "list" in nombre or "bullet" in nombre or "numbered" in nombre:
        print(f"  {s.name}")
